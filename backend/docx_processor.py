"""
DOCX Processor — Fixed for 3 real-world bugs found in testing
=============================================================

BUG 1 FIXED — Table cell skipping due to python-docx vMerge phantom cells:
  python-docx's row.cells includes vertically-merged cells repeated across rows.
  These repeated cells share the SAME _tc object id, so our seen-set was
  skipping real untranslated cells.
  FIX: Iterate raw <w:tc> XML elements per row — these are always unique per row.

BUG 2 FIXED — Hyperlink paragraphs have zero runs:
  Text inside <w:hyperlink> is invisible to para.runs (runs only sees direct
  <w:r> children of <w:p>). So hyperlink text was silently skipped.
  FIX: Extract text from ALL <w:t> descendants of <w:p> via XML iteration.
  Write translated text back via the same XML approach.

BUG 3 FIXED — Duplicate cells in CSV causing partial translation:
  (In csv_processor.py — see that file)
"""
import time
import copy
import re
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W   = lambda tag: f'{{{WNS}}}{tag}'


# ── Text extraction ─────────────────────────────────────────────────────────

def _xml_text(elem) -> str:
    """Extract ALL text from an XML element including hyperlinks, fields, etc."""
    return ''.join(t.text or '' for t in elem.iter(W('t')))


def _para_text(para) -> str:
    """Get full paragraph text including hyperlinks (not just runs)."""
    return _xml_text(para._p)


def _is_skippable(text: str) -> bool:
    t = text.strip()
    if not t:
        return True
    # Pure numbers / dates / percentages — never need translation
    if re.match(r'^[\d,.\-+\s%$€£¥₹/()]+$', t):
        return True
    # URLs — skip
    if re.match(r'^https?://', t):
        return True
    # Very short pure-uppercase codes: ID, N/A, USD etc.
    if re.match(r'^[A-Z0-9_\-]{1,4}$', t):
        return True
    return False


# ── Text write-back ──────────────────────────────────────────────────────────

def _set_para_text(para, translated: str):
    """
    Write translated text back into paragraph.
    Handles 3 cases:
      A) Normal runs — set first run's text, clear rest
      B) Hyperlink paragraph — replace ALL <w:t> text nodes directly in XML
      C) No runs at all — append a new run
    """
    p_elem = para._p

    # Collect all <w:t> elements anywhere in this paragraph
    all_t = list(p_elem.iter(W('t')))

    if not all_t:
        # Case C: truly empty paragraph — add a run
        para.add_run(translated)
        return

    # Case B / A: put translated text into first <w:t>, clear the rest
    all_t[0].text = translated
    for t in all_t[1:]:
        t.text = ''


def _clone_run_fmt(src_run, dst_run):
    src_rpr = src_run._r.find(W('rPr'))
    if src_rpr is not None:
        dst_rpr = dst_run._r.find(W('rPr'))
        if dst_rpr is not None:
            dst_run._r.remove(dst_rpr)
        dst_run._r.insert(0, copy.deepcopy(src_rpr))


# ── Translation ───────────────────────────────────────────────────────────────

def _translate_text(text: str, src: str, tgt: str) -> str:
    from backend.translator import translate_paragraph
    if not text or not text.strip():
        return text
    res = translate_paragraph(text.strip(), src, tgt)
    return res['output'] if (res.get('ok') or res.get('output')) else text


# ── Paragraph handler ─────────────────────────────────────────────────────────

def _do_para(para, src, tgt, counter, total, progress_cb):
    text = _para_text(para).strip()
    if not text or _is_skippable(text):
        return
    translated = _translate_text(text, src, tgt)
    _set_para_text(para, translated)
    counter[0] += 1
    if progress_cb:
        progress_cb(counter[0], total[0], text[:60])
    time.sleep(0.10)


# ── Table handler — BUG 1 FIX ─────────────────────────────────────────────────

def _iter_real_cells(table):
    """
    Iterate real (non-phantom) table cells by walking raw <w:tc> XML elements.

    python-docx's row.cells repeats vertically-merged cells across rows using
    vMerge, which means the same <w:tc> object appears multiple times and our
    seen-set incorrectly skips them.

    Walking <w:tr>/<w:tc> directly gives exactly the cells physically present
    in each row — no phantoms, no repeats.
    """
    for tr in table._tbl.iterchildren(W('tr')):
        for tc in tr.iterchildren(W('tc')):
            yield tc


def _do_table(table, src, tgt, counter, total, progress_cb):
    """Translate all cells in a table using raw XML iteration (bug-1 fix)."""
    from docx.table import _Cell
    for tc in _iter_real_cells(table):
        cell = _Cell(tc, table)
        for para in cell.paragraphs:
            _do_para(para, src, tgt, counter, total, progress_cb)
        # Nested tables
        for nested in cell.tables:
            _do_table(nested, src, tgt, counter, total, progress_cb)


# ── Text boxes ────────────────────────────────────────────────────────────────

def _do_textboxes(doc, src, tgt, counter, total, progress_cb):
    body = doc.element.body
    WPS  = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    for txbx in body.iter(f'{{{WPS}}}txbx'):
        for p_elem in txbx.iter(W('p')):
            from docx.text.paragraph import Paragraph as DocxPara
            try:
                para = DocxPara(p_elem, None)
                _do_para(para, src, tgt, counter, total, progress_cb)
            except Exception:
                pass


# ── Count ─────────────────────────────────────────────────────────────────────

def _count(doc) -> int:
    from docx.table import _Cell
    n = 0
    for para in doc.paragraphs:
        if _para_text(para).strip() and not _is_skippable(_para_text(para)):
            n += 1
    for tbl in doc.tables:
        for tc in _iter_real_cells(tbl):
            cell = _Cell(tc, tbl)
            for p in cell.paragraphs:
                if _para_text(p).strip() and not _is_skippable(_para_text(p)):
                    n += 1
    return max(n, 1)


# ── Public API ────────────────────────────────────────────────────────────────

def translate_docx(input_path: str, output_path: str,
                   src: str, tgt: str, progress_cb=None) -> dict:
    doc     = Document(input_path)
    total   = [_count(doc)]
    counter = [0]

    # Body paragraphs
    for para in doc.paragraphs:
        _do_para(para, src, tgt, counter, total, progress_cb)

    # Tables — using raw XML iteration (Bug 1 fix)
    for tbl in doc.tables:
        _do_table(tbl, src, tgt, counter, total, progress_cb)

    # Headers / footers
    for section in doc.sections:
        for hdr_ftr in [
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ]:
            try:
                for para in hdr_ftr.paragraphs:
                    _do_para(para, src, tgt, counter, total, progress_cb)
                for tbl in hdr_ftr.tables:
                    _do_table(tbl, src, tgt, counter, total, progress_cb)
            except Exception:
                pass

    # Text boxes
    try:
        _do_textboxes(doc, src, tgt, counter, total, progress_cb)
    except Exception:
        pass

    doc.save(output_path)
    return {'ok': True, 'total': total[0], 'done': counter[0]}